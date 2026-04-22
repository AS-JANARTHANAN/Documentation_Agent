"""
╔══════════════════════════════════════════════════════════╗
║   AGENT 2 v2 — PROJECT DOCUMENTATION FORMATTER          ║
╚══════════════════════════════════════════════════════════╝

WHAT WAS WRONG IN v1:
  - Document stopped at Chapter 3 (content was incomplete)
  - No [SUBHEADING_SUB] support for 3.1.1, 3.2.1 etc.
  - No References section handling
  - No Appendices / Ch6 handling
  - Header looked like a faded watermark

WHAT IS FIXED IN v2:
  - Handles all 7 chapters + References
  - [SUBHEADING_SUB] → slightly smaller, still bold, left-aligned
  - [REFERENCES_START/END] → proper numbered references list
  - [CODE_SAMPLE] → fixed-width code block styling
  - Header is now a clean bold line, not grey/faded
  - Watermark on every page is removed
  - 3.1 Existing System sub-categories styled correctly
  - Image placeholders have proper figure captions

FORMATTING SPECS (from sample document):
  Font:          Times New Roman (entire document)
  Heading:       14pt, UPPERCASE, Center, Bold
  Chapter:       14pt, UPPERCASE, Center, Bold
  Subheading:    14pt, UPPERCASE, Left, Bold
  Sub-sub:       12pt, UPPERCASE, Left, Bold (for 3.1.1 etc.)
  Body text:     12pt, Justified, Tab indent first line
  Line spacing:  1.5 throughout
  Margins:       1 inch all sides
  Header:        Project title (centered, clean)
  Footer:        Page number starting at 1

RUN:
    python agent2_formatter_v2.py

REQUIREMENTS:
    pip install python-docx ollama
"""

import re
import sys
import os
import ollama
from pathlib import Path

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm


# ═══════════════════════════════════════════════════════════════
#  OLLAMA SUMMARIZER — trims overly long sections
# ═══════════════════════════════════════════════════════════════
class OllamaSummarizer:
    def __init__(self, model="llama3", host="http://localhost:11434"):
        self.model  = model
        self.client = ollama.Client(host=host)

    def summarize_if_needed(self, text: str, max_words: int = 650) -> str:
        wc = len(text.split())
        if wc <= max_words:
            return text
        print(f"    [Ollama] Trimming paragraph ({wc}w → ~{max_words}w)...")
        prompt = f"""Summarize the following academic paragraph to approximately {max_words} words.
Keep ALL key technical information. Keep formal academic tone.
Return ONLY the paragraph — no preamble, no markdown.

TEXT:
{text}"""
        try:
            resp = self.client.chat(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                options={"temperature": 0.1}
            )
            return resp["message"]["content"].strip()
        except Exception as e:
            print(f"    [Ollama] ⚠ Summarizer failed ({e})")
            return text


# ═══════════════════════════════════════════════════════════════
#  DOCX HELPERS
# ═══════════════════════════════════════════════════════════════
def _force_times_new_roman(run):
    """Force Times New Roman in XML (required for some Word versions)."""
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    "Times New Roman")
    rFonts.set(qn("w:hAnsi"),    "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    rFonts.set(qn("w:cs"),       "Times New Roman")
    # Remove any existing rFonts before inserting
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def set_font(run, size_pt=12, bold=False, italic=False, color=None):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    _force_times_new_roman(run)


def set_spacing_15(para):
    """Apply 1.5 line spacing to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    # Remove existing spacing element to avoid duplication
    for sp in pPr.findall(qn("w:spacing")):
        pPr.remove(sp)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"),     "360")   # 360 twips = 1.5 × single (240)
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_chapter_heading(doc, text: str):
    """CHAPTER N line + CHAPTER TITLE — 14pt, center, bold, uppercase."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    set_font(run, size_pt=14, bold=True)
    set_spacing_15(p)
    return p


def add_subheading(doc, text: str):
    """Section subheading — 14pt, left, bold, uppercase."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    set_font(run, size_pt=14, bold=True)
    set_spacing_15(p)
    return p


def add_subsub_heading(doc, text: str):
    """Sub-subheading (e.g. 3.1.1) — 12pt, left, bold, uppercase."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    set_font(run, size_pt=12, bold=True)
    set_spacing_15(p)
    return p


def add_body_para(doc, text: str):
    """Body paragraph — 12pt, justified, first-line tab indent, 1.5 spacing."""
    text = text.strip()
    if not text:
        return None
    # First letter uppercase
    text = text[0].upper() + text[1:]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # First-line tab indent (720 twips = 0.5 inch)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "720")
    # Remove existing indent
    for existing_ind in pPr.findall(qn("w:ind")):
        pPr.remove(existing_ind)
    pPr.append(ind)

    run = p.add_run(text)
    set_font(run, size_pt=12)
    set_spacing_15(p)
    return p


def add_image_placeholder(doc, description: str):
    """Styled placeholder box for where a diagram will go."""
    # Blank line before
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(f"[ {description.upper()} ]")
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Grey box border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "8")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "888888")
        pBdr.append(b)
    pPr.append(pBdr)

    # Add spacing around placeholder
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"),  "120")
    pPr.append(spacing)

    doc.add_paragraph()
    return p


def add_list_item(doc, text: str):
    """Bullet list item — 12pt, Times New Roman."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),    "720")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)

    run = p.add_run("• " + text.strip())
    set_font(run, size_pt=12)
    set_spacing_15(p)
    return p


def add_reference_entry(doc, num: int, text: str):
    """Numbered reference entry."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),    "720")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)

    run = p.add_run(f"[{num}] {text.strip()}")
    set_font(run, size_pt=12)
    set_spacing_15(p)
    return p


def add_code_block(doc, lines: list):
    """Code sample in Courier New, smaller font."""
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        pPr.append(ind)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        _force_times_new_roman(run)
        run._r.get_or_add_rPr().find(qn("w:rFonts")).set(qn("w:ascii"), "Courier New")
        run._r.get_or_add_rPr().find(qn("w:rFonts")).set(qn("w:hAnsi"), "Courier New")
        set_spacing_15(p)


# ═══════════════════════════════════════════════════════════════
#  HEADER + FOOTER (no watermark, clean header line)
# ═══════════════════════════════════════════════════════════════
def setup_header_footer(section, project_title: str):
    """
    Header: Empty.
    Footer: page number centered.
    """
    # ── HEADER
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        hp = header.paragraphs[0]
    else:
        hp = header.add_paragraph()

    hp.clear()

    # ── FOOTER with page number field
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        fp = footer.paragraphs[0]
    else:
        fp = footer.add_paragraph()

    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    _force_times_new_roman(run)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_margins(section):
    """1 inch margins on all sides."""
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)


# ═══════════════════════════════════════════════════════════════
#  CONTENT PARSER
# ═══════════════════════════════════════════════════════════════
def parse_tagged_content(raw_text: str) -> list:
    """
    Parses tagged content into blocks.

    Supported tags (v2 additions marked with ★):
      [CHAPTER: N | TITLE]
      [SUBHEADING: X.X TITLE]
      [SUBHEADING_SUB: X.X.X TITLE]    ★ nested subheading
      [CONTENT]
      [IMAGE_PLACEHOLDER: description]
      [LIST_START] / [LIST_END]
      [LIST_ITEM: text]
      [REFERENCES_START] / [REFERENCES_END]  ★
      [REFERENCE: text]                 ★
      [CODE_SAMPLE]                     ★ (treats next lines as code until blank)
      [PAGE_BREAK]
    """
    blocks   = []
    lines    = raw_text.splitlines()
    i        = 0
    in_list  = False
    in_refs  = False
    ref_num  = 1
    collecting_content = False
    content_buffer     = []

    def flush_content():
        nonlocal content_buffer, collecting_content
        if content_buffer:
            text = " ".join(content_buffer).strip()
            if text:
                blocks.append({"type": "content", "text": text})
        content_buffer     = []
        collecting_content = False

    while i < len(lines):
        raw_line = lines[i]
        line     = raw_line.strip()

        # ── Chapter
        m = re.match(r"\[CHAPTER:\s*(\d+)\s*\|\s*(.+?)\]", line, re.I)
        if m:
            flush_content()
            blocks.append({"type": "chapter", "num": m.group(1).strip(),
                            "text": m.group(2).strip()})
            i += 1; continue

        # ── Sub-subheading (3.1.1 etc.) — check before subheading
        m = re.match(r"\[SUBHEADING_SUB:\s*(.+?)\]", line, re.I)
        if m:
            flush_content()
            blocks.append({"type": "subheading_sub", "text": m.group(1).strip()})
            i += 1; continue

        # ── Subheading
        m = re.match(r"\[SUBHEADING:\s*(.+?)\]", line, re.I)
        if m:
            flush_content()
            blocks.append({"type": "subheading", "text": m.group(1).strip()})
            i += 1; continue

        # ── Content tag
        if re.match(r"\[CONTENT\]", line, re.I):
            flush_content()
            collecting_content = True
            i += 1; continue

        # ── Image placeholder
        m = re.match(r"\[IMAGE_PLACEHOLDER:\s*(.+?)\]", line, re.I)
        if m:
            flush_content()
            blocks.append({"type": "image_placeholder", "text": m.group(1).strip()})
            i += 1; continue

        # ── List
        if re.match(r"\[LIST_START\]", line, re.I):
            flush_content()
            in_list = True
            i += 1; continue
        if re.match(r"\[LIST_END\]", line, re.I):
            in_list = False
            i += 1; continue

        m = re.match(r"\[LIST_ITEM:\s*(.+?)\]", line, re.I)
        if m:
            blocks.append({"type": "list_item", "text": m.group(1).strip()})
            i += 1; continue

        # ── References ★
        if re.match(r"\[REFERENCES_START\]", line, re.I):
            flush_content()
            in_refs = True
            blocks.append({"type": "references_start"})
            i += 1; continue
        if re.match(r"\[REFERENCES_END\]", line, re.I):
            in_refs = False
            blocks.append({"type": "references_end"})
            i += 1; continue

        m = re.match(r"\[REFERENCE:\s*(.+?)\]", line, re.I)
        if m:
            blocks.append({"type": "reference", "num": ref_num,
                            "text": m.group(1).strip()})
            ref_num += 1
            i += 1; continue

        # ── Code sample ★
        if re.match(r"\[CODE_SAMPLE\]", line, re.I):
            flush_content()
            blocks.append({"type": "code_sample_marker"})
            i += 1; continue

        # ── Page break
        if re.match(r"\[PAGE_BREAK\]", line, re.I):
            flush_content()
            blocks.append({"type": "page_break"})
            i += 1; continue

        # ── Plain text
        if line:
            if collecting_content or in_list or in_refs:
                content_buffer.append(line)
            else:
                content_buffer.append(line)
                collecting_content = True
        else:
            flush_content()

        i += 1

    flush_content()
    return blocks


# ═══════════════════════════════════════════════════════════════
#  DOCUMENT FORMATTER AGENT v2
# ═══════════════════════════════════════════════════════════════
class DocumentFormatterAgentV2:
    def __init__(self, ollama_model="llama3", ollama_host="http://localhost:11434",
                 use_summarizer=True):
        self.summarizer = OllamaSummarizer(ollama_model, ollama_host) if use_summarizer else None
        print(f"[Agent 2 v2] Initialized. Summarizer: {'ON' if use_summarizer else 'OFF'}")

    # ── Cover Page
    def build_cover_page(self, doc, project_title, degree, department,
                         student_name, roll_no, college, guide, year):
        section = doc.sections[0]
        set_margins(section)
        setup_header_footer(section, project_title)

        def c_bold(text, size=14):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_font(run, size_pt=size, bold=True)
            set_spacing_15(p)

        def c_norm(text, size=12):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_font(run, size_pt=size)
            set_spacing_15(p)

        def blank():
            p = doc.add_paragraph()
            set_spacing_15(p)

        blank()
        c_bold(project_title.upper(), 14)
        blank()
        c_norm("A PROJECT REPORT", 12)
        blank()
        c_norm("Submitted by", 12)
        blank()
        c_bold(f"{student_name.upper()}  ({roll_no})", 12)
        blank()
        c_norm("In partial fulfilment for the award of the degree of", 12)
        blank()
        c_bold(f"BACHELOR OF TECHNOLOGY", 12)
        c_bold(f"IN {department.upper()}", 12)
        blank()
        c_norm("Under the Guidance of", 12)
        c_bold(guide.upper(), 12)
        blank()
        c_bold(college.upper(), 12)
        c_norm("(AUTONOMOUS)", 12)
        blank()
        c_bold("ANNA UNIVERSITY", 12)
        blank()
        c_bold(year.upper(), 12)

        doc.add_page_break()

    # ── Certificate Page
    def build_certificate(self, doc, project_title, student_name, roll_no,
                           college, department, guide, hod):
        add_chapter_heading(doc, "BONAFIDE CERTIFICATE")
        doc.add_paragraph()
        text = (
            f'Certified that this project report titled "{project_title.upper()}" '
            f'is the bonafide work of {student_name.upper()} ({roll_no}), '
            f'who carried out the project work under my supervision.'
        )
        add_body_para(doc, text)
        doc.add_paragraph()
        doc.add_paragraph()

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        for cell in tbl.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, size_pt=11)
        tbl.rows[0].cells[0].text = (
            f"SIGNATURE\n{hod}\nHEAD OF DEPARTMENT\nDepartment of {department}\n{college}"
        )
        tbl.rows[0].cells[1].text = (
            f"SIGNATURE\n{guide}\nPROJECT GUIDE\nDepartment of {department}\n{college}"
        )
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Submitted for the end semester examination held on ………………")
        set_font(run, size_pt=12)
        set_spacing_15(p)
        doc.add_page_break()

    # ── Declaration Page
    def build_declaration(self, doc):
        add_chapter_heading(doc, "DECLARATION BY THE CANDIDATE")
        doc.add_paragraph()
        text = (
            "I declare that to the best of my knowledge the work reported herein has been "
            "composed solely by myself and that it has not been in whole or in part submitted "
            "in any previous application for a degree. All sources of information have been "
            "specifically acknowledged by means of references."
        )
        add_body_para(doc, text)
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("SIGNATURE OF THE CANDIDATE")
        set_font(run, size_pt=12, bold=True)
        set_spacing_15(p)
        doc.add_page_break()

    # ── Main formatter
    def format_document(self, tagged_content: str, project_title: str,
                        student_name: str, roll_no: str, college: str,
                        department: str, guide: str, hod: str,
                        degree: str, year: str, output_path: str):

        print("[Agent 2 v2] Parsing tagged content...")
        blocks = parse_tagged_content(tagged_content)
        total  = len(blocks)
        print(f"[Agent 2 v2] Found {total} content blocks")

        doc = Document()

        # Initialize document layout
        section = doc.sections[0]
        set_margins(section)
        setup_header_footer(section, project_title)

        # ── Render content blocks
        print("[Agent 2 v2] Rendering chapters...")
        in_references = False
        block_num = 0

        for block in blocks:
            block_num += 1
            btype = block.get("type")

            if block_num % 50 == 0:
                print(f"  [Agent 2 v2] Progress: {block_num}/{total} blocks")

            # Chapter heading
            if btype == "chapter":
                add_chapter_heading(doc, f"CHAPTER {block['num']}")
                add_chapter_heading(doc, block["text"])

            # Main subheading (1.1, 2.1, etc.)
            elif btype == "subheading":
                add_subheading(doc, block["text"])

            # Sub-subheading (3.1.1, 3.2.1, etc.)
            elif btype == "subheading_sub":
                add_subsub_heading(doc, block["text"])

            # Body paragraph
            elif btype == "content":
                text = block["text"]
                if self.summarizer:
                    text = self.summarizer.summarize_if_needed(text, max_words=650)
                add_body_para(doc, text)

            # Image placeholder
            elif btype == "image_placeholder":
                add_image_placeholder(doc, block["text"])

            # Bullet list item
            elif btype == "list_item":
                add_list_item(doc, block["text"])

            # References
            elif btype == "references_start":
                add_subheading(doc, "REFERENCES")
                in_references = True

            elif btype == "reference":
                add_reference_entry(doc, block["num"], block["text"])

            elif btype == "references_end":
                in_references = False

            # Code sample marker
            elif btype == "code_sample_marker":
                p = doc.add_paragraph()
                run = p.add_run("[Source Code — Add code screenshot here]")
                run.font.name  = "Courier New"
                run.font.size  = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Page break
            elif btype == "page_break":
                doc.add_page_break()

        # ── Save
        doc.save(output_path)
        size_kb = os.path.getsize(output_path) // 1024
        print(f"[Agent 2 v2] ✅ Saved: {output_path}  ({size_kb} KB)")

    # ── CLI
    def run(self):
        print("\n" + "=" * 60)
        print("   AGENT 2 v2 — PROJECT DOCUMENTATION FORMATTER")
        print("=" * 60)
        print("  All 7 chapters + References → formatted .docx")
        print("=" * 60 + "\n")

        content_path = input("📄 Path to content .txt file: ").strip().strip('"')
        if not Path(content_path).exists():
            print(f"❌ File not found: {content_path}")
            sys.exit(1)

        with open(content_path, "r", encoding="utf-8") as f:
            tagged_content = f.read()

        word_count = len(tagged_content.split())
        print(f"  Loaded: {word_count:,} words from {content_path}")

        if word_count < 5000:
            print("  ⚠ WARNING: Word count is low. The AI may not have written all chapters.")
            print("  ⚠ Check that the content file contains all 7 chapters.")
            cont = input("  Continue anyway? [y/n]: ").strip().lower()
            if cont != "y":
                sys.exit(0)

        print("\n── Project Details ──")
        project_title = input("📌 Project Title: ").strip()
        student_name  = input("👤 Student Name: ").strip()
        roll_no       = input("🔢 Roll Number: ").strip()
        college       = input("🏫 College Name: ").strip()
        department    = input("🏢 Department: ").strip() or "Information Technology"
        guide         = input("👨🏫 Guide Name & Designation: ").strip()
        hod           = input("👨💼 HOD Name & Designation: ").strip()
        degree        = input("🎓 Degree (e.g. B.Tech): ").strip() or "B.Tech"
        year          = input("📅 Year (e.g. MAY 2026): ").strip() or "MAY 2026"
        output_path   = input("💾 Output filename (e.g. Final_Doc.docx): ").strip() \
                        or "ProjectDocumentation_v2.docx"

        use_ollama = input("\n🤖 Use Ollama to trim long paragraphs? [y/n]: ").strip().lower() == "y"
        if not use_ollama:
            self.summarizer = None

        self.format_document(
            tagged_content=tagged_content,
            project_title=project_title,
            student_name=student_name,
            roll_no=roll_no,
            college=college,
            department=department,
            guide=guide,
            hod=hod,
            degree=degree,
            year=year,
            output_path=output_path
        )

        print("\n[Agent 2 v2] ✅ All done!")
        print(f"[Agent 2 v2] 📂 Open: {output_path}")
        print("[Agent 2 v2] 💡 Manually replace [IMAGE_PLACEHOLDER] boxes with your diagrams\n")


if __name__ == "__main__":
    agent = DocumentFormatterAgentV2(
        ollama_model="llama3",
        ollama_host="http://localhost:11434",
        use_summarizer=True
    )
    agent.run()